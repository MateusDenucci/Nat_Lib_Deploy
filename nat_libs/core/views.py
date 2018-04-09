from django.shortcuts import render
from django.http import HttpResponseRedirect,HttpResponse
import re
from docx import Document
from docx.shared import Inches
from django.urls import reverse


input_text = ""

# Create your views here.
def index(request):
	return render(request,'core/index.html')

def process_input(request):
	global input_text

	input_text = request.POST.get('text',None)
	language = request.POST.get('language',None)
	
	request.session['language'] = language
	return HttpResponseRedirect('/core/questions')

def questions(request):

	pos = ['*VERB', '*VERB_INFINITIVE', '*VERB_PAST_SINGULAR', '*VERB_PAST_PLURAL',
	       '*VERB_PRESENT_SINGULAR', '*VERB_PRESENT_PLURAL', '*VERB_FUTURE_SINGULAR',
	       '*VERB_FUTURE_PLURAL', '*VERB_PAST_PARTICIPLE', '*VERB_GERUND', '*NOUN',
	       '*NOUN_PLURAL', '*NOUN_SINGULAR_MASCULINE', '*NOUN_SINGULAR_FEMININE',
	       '*NOUN_PLURAL_MASCULINE', '*NOUN_PLURAL_FEMININE',
	       '*ADJECTIVE', '*ADJECTIVE_SINGULAR_MASCULINE',
	       '*ADJECTIVE_PLURAL_MASCULINE', '*ADJECTIVE_SINGULAR_FEMININE',
	       '*ADJECTIVE_PLURAL_FEMININE', '*ADVERB', '*BODY_PART', '*PROPER_NOUN',
	       '*NAME_OF_PERSON', '*LOCATION',  '*NUMBER', '*TIME']

	questions = language_set(request.session['language'])

	instances = find_instances()
	questions_to_user = get_questions(instances, pos,questions)
	dic={}
	for i in range(len(questions_to_user)):
		dic[i] = questions_to_user[i]

	request.session['instances'] = instances
	#request.session['natlib_regex'] = natlib_regex
	return render(request,'core/questions.html',{"questions_to_user": dic})

def find_instances():
    global input_text

    natlib_regex = re.compile(r'\*[A-Z][A-Z][A-Z][A-Z_]+')
    instances = []
    for hits in natlib_regex.findall(input_text):
        instances.append(hits)
    return instances



def get_questions(instances, pos,questions):
	print()
	questions_to_user = []
	for i in range(len(instances)):
		for j in range(len(pos)):
			if instances[i] == pos[j]:
				questions_to_user.append(questions[j])
	return questions_to_user


def process_questions(request):
	global input_text
	answers=[]

	for i in range(len(request.GET)):
		answers.append(request.GET.get(str(i),None))
	# return HttpResponse(len(request.GET))

	natlib_regex = re.compile(r'\*[A-Z][A-Z][A-Z][A-Z_]+')
	instances = request.session['instances']
	user_input = answers
	for i in range(len(instances)):
		input_text = natlib_regex.sub(user_input[i], input_text, 1)

	return HttpResponseRedirect('/core/result')

def result(request):
	return render(request,'core/result.html',{"text": input_text})

def try_again(request):
	request.session.flush()
	return HttpResponseRedirect('/core/')

def generate_doc(request):
	
	document = Document()

	document.add_heading('Nat Libs',0)

	document.add_paragraph(input_text)


	response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
	response['Content-Disposition'] = 'attachment; filename="Nat_Libs.docx"'
	document.save(response)
	return response

def language_set(language):
    english = ['Enter a verb:', 'Enter a verb in the infinitive:',
               'Enter a verb in the past tense:', 'Enter a verb in the past tense:',
               'Enter a verb:', 'Enter a verb:',
               'Enter a verb in the future tense:', 'Enter a verb in the future tense:',
               'Enter the past participle of a verb:', 'Enter a verb in the gerund:',
               'Enter a noun:', 'Enter a plural noun:',
               'Enter a noun:', 'Enter a noun:',
               'Enter a plural noun:', 'Enter a plural noun:',
               'Enter an adjective:', 'Enter an adjective:',
               'Enter an adjective:', 'Enter an adjective:',
               'Enter an adjective:', 'Enter an adverb:',
               'Enter a body part:', 'Enter a proper noun:',
               'Enter the name of a person:', 'Enter a location:',
               'Enter a number:', 'Enter a point in time.']
    spanish = ['Escriba un verbo:', 'Escriba un verbo en la forma infinitiva:',
               'Escriba un verbo en el singular del pretérito:',
               'Escriba un verbo en el plural del pretérito:',
               'Escriba un verbo en el singular del presente:',
               'Escriba un verbo en el plural del presente:',
               'Escriba un verbo en el singular del futuro:',
               'Escriba un verbo en el plural del futuro:',
               'Escriba un verbo en el participio pasado:',
               'Escriba un verbo en el gerundio:',
               'Escriba un substantivo:', 'Escriba un substantivo en el plural:',
               'Escriba un substantivo masculino:', 'Escriba un substantivo feminino:',
               'Escriba un substantivo masculino plural:',
               'Escriba un substantivo feminino plural:',
               'Escriba un adjetivo:', 'Escriba un adjectivo masculino singular:',
               'Escriba un adjetivo masculino plural:', 'Escriba un adjetivo feminino singular:',
               'Escriba un adjetivo feminino plural:', 'Escriba un adverbio:',
               'Escriba una parte del cuerpo:', 'Escriba un nombre propio:',
               'Escriba el nombre de una persona:', 'Escriba un lugar:',
               'Escriba un número:', 'Escriba un punto en el tiempo:']
    portuguese = ['Digite um verbo:', 'Digite um verbo na forma infinitiva:',
                  'Digite um verbo no singular do pretérito:',
                  'Digite um verbo no plural do pretérito:',
                  'Digite um verbo no singular do presente:',
                  'Digite um verbo no plural do presente:',
                  'Digite um verbo no singular do futuro:',
                  'Digite um verbo no plural do futuro:',
                  'Digite um verbo no particípio do passado:',
                  'Digite um verbo no gerúndio:',
                  'Digite um substantivo:', 'Digite um substantivo no plural:',
                  'Digite um substantivo masculino:', 'Digite um substantivo feminino:',
                  'Digite um substantivo masculino no plural:',
                  'Digite um substantivo feminino no plural:',
                  'Digite um adjetivo:', 'Digite um adjetivo masculino no singular:',
                  'Digite um adjetivo masculino no plural:', 'Digite um adjetivo feminino no singular:',
                  'Digite um adjetivo feminino no plural:', 'Digite um advérbio:',
                  'Digite uma parte do corpo:', 'Digite um nome próprio:',
                  'Digite o nome de uma pessoa:', 'Digite um lugar:',
                  'Digite um número:', 'Digite um tempo (hora, dia, ano etc.):']
    if language == 'English':
        return english
    if language == 'Spanish':
        return spanish
    if language == 'French':
        return french
    if language == 'Portuguese':
        return portuguese